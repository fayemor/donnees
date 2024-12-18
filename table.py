import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx import Document
import os

def save_to_word(data, filename):
    doc = Document()
    doc.add_heading("Informations Collectées", level=1)

    table = doc.add_table(rows=1, cols=len(data[0]))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(data[0]):
        hdr_cells[idx].text = header

    for row in data[1:]:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            row_cells[idx].text = value

    # enreigistrement des documents
    doc.save(filename)
    messagebox.showinfo("Succès", f"Le fichier a été enregistré sous {filename}")

# Fonction pour ajouter les donnes dans le tableau
def add_data():
    values = [entry.get() for entry in entries]
    if any(v.strip() == "" for v in values):
        messagebox.showerror("Erreur", "Tous les champs doivent être remplis")
        return

    table_data.append(values)
    for idx, val in enumerate(values):
        table.insert("", "end", values=values)

    for entry in entries:
        entry.delete(0, tk.END)

# Fontion d'exportation des donnes
def export_data():
    if not table_data:
        messagebox.showerror("Erreur", "Aucune donnée à exporter")
        return

    filename = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")]
    )
    if filename:
        save_to_word([headers] + table_data, filename)

app = tk.Tk()
app.title("Insertion de donnees")
app.geometry("800x500")

headers = ["Nom", "Prénom", "Age", "Email"]

table_data = []

entry_frame = tk.Frame(app)
entry_frame.pack(pady=10)

entries = []
for idx, header in enumerate(headers):
    lbl = tk.Label(entry_frame, text=header)
    lbl.grid(row=0, column=idx, padx=10, pady=5)
    entry = tk.Entry(entry_frame)
    entry.grid(row=1, column=idx, padx=10, pady=5)
    entries.append(entry)

add_btn = tk.Button(app, text="Ajouter", command=add_data, bg="blue", fg="white")
add_btn.pack(pady=10)


table = ttk.Treeview(app, columns=headers, show="headings")
for header in headers:
    table.heading(header, text=header)
    table.column(header, width=150)
table.pack(pady=10, fill=tk.BOTH, expand=True)

export_btn = tk.Button(app, text="Exporter en Word", command=export_data, bg="blue", fg="white")
export_btn.pack(pady=10)

app.mainloop()
